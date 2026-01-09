from .base import BaseConverter
from docx import Document
from docx.shared import Pt

class DOCXConverter(BaseConverter):
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['docx']
    
    def convert(self, input_text, output_path, options=None):
        self.validate_input(input_text)
        self.ensure_output_dir(output_path)
        
        doc = Document()
        
        lines = input_text.split('\n')
        for line in lines:
            if line.strip():
                para = doc.add_paragraph(line)
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
            else:
                doc.add_paragraph()
        
        doc.save(output_path)
        return output_path
