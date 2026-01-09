import os
from .parsers import get_parser

def get_file_metadata(file_path, file_extension):
    metadata = {
        'size': os.path.getsize(file_path),
        'extension': file_extension,
        'valid': False,
        'error': None
    }
    
    try:
        parser = get_parser(file_extension)
        content = parser.parse(file_path)
        
        metadata['valid'] = True
        metadata['character_count'] = len(content)
        metadata['line_count'] = len(content.split('\n'))
        
        if file_extension == 'pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    metadata['page_count'] = len(pdf_reader.pages)
            except:
                pass
        
        elif file_extension == 'docx':
            try:
                from docx import Document
                doc = Document(file_path)
                metadata['paragraph_count'] = len(doc.paragraphs)
                word_count = sum(len(para.text.split()) for para in doc.paragraphs)
                metadata['word_count'] = word_count
            except:
                pass
        
        elif file_extension == 'html':
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                    metadata['tag_count'] = len(soup.find_all())
            except:
                pass
        
        elif file_extension == 'json':
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict):
                    metadata['key_count'] = len(data)
                elif isinstance(data, list):
                    metadata['item_count'] = len(data)
        
        elif file_extension == 'csv':
            import csv
            with open(file_path, 'r', encoding='utf-8', newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)
                metadata['row_count'] = len(rows)
                if rows:
                    metadata['column_count'] = len(rows[0])
        
        elif file_extension == 'xml':
            import xml.etree.ElementTree as ET
            tree = ET.parse(file_path)
            root = tree.getroot()
            metadata['element_count'] = len(list(root.iter()))
        
    except Exception as e:
        metadata['error'] = str(e)
        metadata['valid'] = False
    
    return metadata
